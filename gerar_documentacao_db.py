"""
Mesa Solidária — Gerador de documentação do banco de dados
Gera: banco_de_dados.xlsx  e  relacionamentos_db.docx
"""
import subprocess, sys

subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl", "python-docx", "-q"])

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ──────────────────────────────────────────────────────────────────────────────
# DEFINIÇÃO DAS TABELAS
# Cada coluna: (nome, tipo, constraints, descrição)
# ──────────────────────────────────────────────────────────────────────────────
TABLES = {
    "usuarios": {
        "desc": "Todos os perfis do sistema em uma única tabela. O campo 'tipo' diferencia o acesso.",
        "cols": [
            ("id_usuario",      "INT",           "PK · AUTO_INCREMENT",           "Identificador único do usuário"),
            ("nome",            "VARCHAR(120)",  "NN",                            "Nome completo"),
            ("email",           "VARCHAR(150)",  "NN · UNIQUE",                   "E-mail de login"),
            ("senha_hash",      "VARCHAR(255)",  "NN",                            "Senha criptografada (bcrypt)"),
            ("cpf",             "VARCHAR(14)",   "UNIQUE",                        "CPF (somente dígitos ou formatado)"),
            ("data_nascimento", "DATE",          "",                              "Data de nascimento"),
            ("telefone",        "VARCHAR(20)",   "",                              "Telefone de contato"),
            ("tipo",            "ENUM",          "NN",                            "'doador' | 'voluntario' | 'beneficiario' | 'ponto_coleta' | 'admin'"),
            ("ativo",           "TINYINT(1)",    "NN · DEFAULT 1",                "1 = ativo, 0 = inativo / bloqueado"),
            ("data_cadastro",   "DATETIME",      "NN · DEFAULT NOW()",            "Data e hora de criação do registro"),
            ("foto_url",        "VARCHAR(255)",  "",                              "URL da foto de perfil (opcional)"),
        ],
    },
    "doadores_perfil": {
        "desc": "Dados complementares do perfil doador. Criada junto com o registro em 'usuarios'.",
        "cols": [
            ("id_perfil",         "INT", "PK · AUTO_INCREMENT", "Identificador do perfil"),
            ("id_usuario",        "INT", "FK → usuarios · NN · UNIQUE", "Vínculo com o usuário"),
            ("id_ponto_preferido","INT", "FK → pontos_coleta",          "Ponto de coleta preferido pelo doador"),
        ],
    },
    "voluntarios_perfil": {
        "desc": "Dados complementares do voluntário, incluindo saldo de pontos acumulados.",
        "cols": [
            ("id_perfil",             "INT",     "PK · AUTO_INCREMENT",    "Identificador do perfil"),
            ("id_usuario",            "INT",     "FK → usuarios · NN · UNIQUE", "Vínculo com o usuário"),
            ("disponibilidade",       "VARCHAR(100)", "",                  "Ex: 'Seg a Sex, manhãs'"),
            ("veiculo",               "VARCHAR(60)",  "",                  "Ex: 'Carro', 'Moto', 'Bicicleta'"),
            ("saldo_pontos",          "INT",     "NN · DEFAULT 0",         "Saldo atual de pontos disponíveis"),
            ("total_pontos_ganhos",   "INT",     "NN · DEFAULT 0",         "Total histórico de pontos ganhos"),
            ("total_pontos_gastos",   "INT",     "NN · DEFAULT 0",         "Total histórico de pontos gastos"),
        ],
    },
    "pontos_coleta": {
        "desc": "Pontos físicos de recebimento de doações e despacho para beneficiários.",
        "cols": [
            ("id_ponto",              "INT",          "PK · AUTO_INCREMENT",     "Identificador do ponto"),
            ("cnpj",                  "CHAR(18)",     "UNIQUE",                  "CNPJ no formato 00.000.000/0000-00"),
            ("nome",                  "VARCHAR(100)", "NN",                      "Nome do ponto (ex: 'Ponto Centro')"),
            ("logradouro",            "VARCHAR(150)", "NN",                      "Rua / Avenida e número"),
            ("complemento",           "VARCHAR(80)",  "",                        "Complemento (sala, bloco, etc.)"),
            ("bairro",                "VARCHAR(80)",  "NN",                      "Bairro"),
            ("cidade",                "VARCHAR(80)",  "NN",                      "Cidade"),
            ("estado",                "CHAR(2)",      "NN",                      "UF (ex: SP)"),
            ("cep",                   "CHAR(9)",      "NN",                      "CEP no formato 00000-000"),
            ("endereco_completo",     "VARCHAR(400)", "",                        "Endereço concatenado gerado automaticamente para geocodificação via API (ex: Google Maps / ViaCEP)"),
            ("horario_funcionamento", "VARCHAR(120)", "",                        "Ex: 'Seg–Sáb 09h–17h'"),
            ("id_responsavel",        "INT",          "FK → usuarios",           "Usuário responsável pelo ponto"),
            ("ativo",                 "TINYINT(1)",   "NN · DEFAULT 1",          "1 = ativo, 0 = desativado"),
        ],
    },
    "familias": {
        "desc": "Dados cadastrais das famílias beneficiárias. Vinculada ao usuário de login do beneficiário.",
        "cols": [
            ("id_familia",            "INT",          "PK · AUTO_INCREMENT",     "Identificador da família"),
            ("id_usuario",            "INT",          "FK → usuarios · NN · UNIQUE", "Vínculo com o usuário beneficiário"),
            ("id_ponto_referencia",   "INT",          "FK → pontos_coleta",      "Ponto de coleta mais próximo"),
            ("logradouro",            "VARCHAR(150)", "NN",                      "Rua / Avenida"),
            ("numero",                "VARCHAR(10)",  "",                        "Número do imóvel"),
            ("complemento",           "VARCHAR(60)",  "",                        "Apto, bloco, etc."),
            ("bairro",                "VARCHAR(80)",  "NN",                      "Bairro"),
            ("cidade",                "VARCHAR(80)",  "NN",                      "Cidade"),
            ("estado",                "CHAR(2)",      "NN",                      "UF"),
            ("cep",                   "CHAR(9)",      "NN",                      "CEP no formato 00000-000"),
            ("endereco_completo",     "VARCHAR(400)", "",                        "Endereço concatenado gerado automaticamente para geocodificação via API (ex: Google Maps / ViaCEP)"),
            ("quantidade_moradores",  "TINYINT",      "NN",                      "Total de pessoas na residência"),
            ("criancas_adolescentes", "TINYINT",      "NN · DEFAULT 0",          "Número de crianças/adolescentes"),
            ("renda_familiar",        "ENUM",         "",                        "'ate_1sm' | '1_a_2sm' | '2_a_3sm' | 'acima_3sm'"),
            ("beneficios_sociais",    "TEXT",         "",                        "Ex: Bolsa Família, BPC"),
            ("restricoes_alimentares","TEXT",         "",                        "Alergias ou restrições alimentares"),
            ("observacoes",           "TEXT",         "",                        "Informações adicionais"),
            ("status",                "ENUM",         "NN · DEFAULT 'pendente'", "'pendente' | 'ativa' | 'inativa'"),
            ("data_cadastro",         "DATETIME",     "NN · DEFAULT NOW()",      "Data de cadastro"),
        ],
    },
    "itens_catalogo": {
        "desc": "Catálogo de tipos de itens doáveis (cesta básica, higiene, roupas, brinquedos…).",
        "cols": [
            ("id_item",  "INT",         "PK · AUTO_INCREMENT", "Identificador do item"),
            ("icone",    "VARCHAR(10)", "",                    "Emoji ou código de ícone"),
            ("nome",     "VARCHAR(80)", "NN · UNIQUE",         "Nome do item (ex: 'Cesta básica')"),
            ("unidade",  "VARCHAR(30)", "NN",                  "Unidade de medida (ex: cestas, kits, peças)"),
            ("categoria","VARCHAR(60)", "NN",                  "Categoria do item"),
            ("ativo",    "TINYINT(1)", "NN · DEFAULT 1",       "1 = disponível para doação"),
        ],
    },
    "estoque_ponto": {
        "desc": "Quantidade atual de cada item em cada ponto de coleta. Atualizada a cada doação recebida ou entrega realizada.",
        "cols": [
            ("id_estoque",       "INT", "PK · AUTO_INCREMENT",           "Identificador do registro"),
            ("id_ponto",         "INT", "FK → pontos_coleta · NN",       "Ponto de coleta"),
            ("id_item",          "INT", "FK → itens_catalogo · NN",      "Tipo do item"),
            ("quantidade",       "INT", "NN · DEFAULT 0",                "Quantidade disponível em estoque"),
            ("meta_quantidade",  "INT", "NN · DEFAULT 0",                "Meta desejada (para calcular necessidade)"),
            ("data_atualizacao", "DATETIME", "NN · DEFAULT NOW()",       "Última atualização do estoque"),
            ("",                 "",    "UNIQUE (id_ponto, id_item)",    "Evita duplicidade de item por ponto"),
        ],
    },
    "doacoes": {
        "desc": "Registro de cada doação efetuada por um doador para um ponto de coleta.",
        "cols": [
            ("id_doacao",        "INT",      "PK · AUTO_INCREMENT",           "Identificador da doação"),
            ("id_doador",        "INT",      "FK → usuarios · NN",            "Doador que registrou"),
            ("id_ponto",         "INT",      "FK → pontos_coleta · NN",       "Ponto que receberá a doação"),
            ("data_registro",    "DATETIME", "NN · DEFAULT NOW()",            "Data em que o doador registrou"),
            ("data_recebimento", "DATETIME", "",                              "Data em que o ponto confirmou o recebimento"),
            ("observacao",       "TEXT",     "",                              "Observações do doador"),
            ("status",           "ENUM",     "NN · DEFAULT 'pendente'",       "'pendente' | 'recebida' | 'cancelada'"),
        ],
    },
    "doacoes_itens": {
        "desc": "Itens que compõem cada doação (tabela de junção N:M entre doacoes e itens_catalogo).",
        "cols": [
            ("id_doacao_item", "INT", "PK · AUTO_INCREMENT",          "Identificador"),
            ("id_doacao",      "INT", "FK → doacoes · NN",            "Doação à qual o item pertence"),
            ("id_item",        "INT", "FK → itens_catalogo · NN",     "Tipo do item doado"),
            ("quantidade",     "INT", "NN",                           "Quantidade deste item na doação"),
        ],
    },
    "solicitacoes": {
        "desc": "Pedidos realizados pelas famílias ao ponto de coleta para receber itens.",
        "cols": [
            ("id_solicitacao",   "INT",      "PK · AUTO_INCREMENT",           "Identificador da solicitação"),
            ("id_familia",       "INT",      "FK → familias · NN",            "Família solicitante"),
            ("id_ponto",         "INT",      "FK → pontos_coleta · NN",       "Ponto de coleta que atenderá"),
            ("modalidade",       "ENUM",     "NN · DEFAULT 'entrega'",        "'entrega' (voluntário leva) | 'retirada' (família busca)"),
            ("observacao",       "TEXT",     "",                              "Observações adicionais"),
            ("data_solicitacao", "DATETIME", "NN · DEFAULT NOW()",            "Data do pedido"),
            ("status",           "ENUM",     "NN · DEFAULT 'pendente'",       "'pendente' | 'em_separacao' | 'aguardando_voluntario' | 'em_entrega' | 'entregue' | 'cancelada'"),
        ],
    },
    "solicitacoes_itens": {
        "desc": "Itens que compõem cada solicitação (tabela de junção N:M entre solicitacoes e itens_catalogo).",
        "cols": [
            ("id_sol_item",    "INT", "PK · AUTO_INCREMENT",         "Identificador"),
            ("id_solicitacao", "INT", "FK → solicitacoes · NN",      "Solicitação à qual o item pertence"),
            ("id_item",        "INT", "FK → itens_catalogo · NN",    "Tipo do item solicitado"),
            ("quantidade",     "INT", "NN",                          "Quantidade solicitada"),
        ],
    },
    "entregas": {
        "desc": "Entrega aceita por um voluntário para uma solicitação. Relacionamento 1:1 com solicitacoes.",
        "cols": [
            ("id_entrega",     "INT",           "PK · AUTO_INCREMENT",           "Identificador da entrega"),
            ("id_solicitacao", "INT",           "FK → solicitacoes · NN · UNIQUE","Solicitação atendida (1 solicitação = 1 entrega)"),
            ("id_voluntario",  "INT",           "FK → usuarios · NN",            "Voluntário responsável"),
            ("id_ponto",       "INT",           "FK → pontos_coleta · NN",       "Ponto onde o voluntário retira os itens"),
            ("distancia_km",   "DECIMAL(6,2)",  "NN",                            "Distância ponto → beneficiário (base do cálculo de pontos)"),
            ("data_aceite",    "DATETIME",      "NN · DEFAULT NOW()",            "Quando o voluntário aceitou a entrega"),
            ("data_retirada",  "DATETIME",      "",                              "Quando confirmou a retirada no ponto"),
            ("data_entrega",   "DATETIME",      "",                              "Quando confirmou a entrega ao beneficiário"),
            ("observacao",     "TEXT",          "",                              "Observações do voluntário"),
            ("status",         "ENUM",          "NN · DEFAULT 'pendente'",       "'pendente' | 'em_rota' | 'entregue' | 'cancelada'"),
        ],
    },
    "pontos_historico": {
        "desc": "Registro de cada ganho ou gasto de pontos de um voluntário. Fonte de verdade do saldo.",
        "cols": [
            ("id_historico",  "INT",          "PK · AUTO_INCREMENT",   "Identificador do lançamento"),
            ("id_voluntario", "INT",          "FK → usuarios · NN",    "Voluntário dono do lançamento"),
            ("tipo",          "ENUM",         "NN",                    "'ganho' | 'gasto'"),
            ("quantidade",    "INT",          "NN",                    "Pontos ganhos ou gastos"),
            ("descricao",     "VARCHAR(200)", "NN",                    "Ex: 'Entrega para Maria Silva (+4 pts)'"),
            ("id_entrega",    "INT",          "FK → entregas",         "Entrega que originou o ganho (se aplicável)"),
            ("id_resgate",    "INT",          "FK → resgates_cupons",  "Resgate que originou o gasto (se aplicável)"),
            ("data_registro", "DATETIME",     "NN · DEFAULT NOW()",    "Data do lançamento"),
        ],
    },
    "cupons": {
        "desc": "Catálogo de recompensas disponíveis para os voluntários resgatarem com seus pontos.",
        "cols": [
            ("id_cupom",      "INT",     "PK · AUTO_INCREMENT", "Identificador do cupom"),
            ("icone",         "VARCHAR(10)", "",                "Emoji representativo"),
            ("categoria",     "VARCHAR(60)", "NN",             "Ex: Cinema, Shows, Tecnologia"),
            ("nome",          "VARCHAR(100)","NN",             "Nome da recompensa"),
            ("descricao",     "TEXT",    "",                   "Descrição do benefício"),
            ("custo_pontos",  "INT",     "NN",                 "Pontos necessários para resgatar"),
            ("parceiro",      "VARCHAR(100)", "",              "Empresa parceira (ex: Cinemark)"),
            ("estoque",       "INT",     "",                   "Quantidade disponível (NULL = ilimitado)"),
            ("ativo",         "TINYINT(1)", "NN · DEFAULT 1",  "1 = disponível para resgate"),
        ],
    },
    "resgates_cupons": {
        "desc": "Cada vez que um voluntário resgata um cupom. Armazena o código único gerado.",
        "cols": [
            ("id_resgate",     "INT",          "PK · AUTO_INCREMENT",   "Identificador do resgate"),
            ("id_voluntario",  "INT",          "FK → usuarios · NN",    "Voluntário que resgatou"),
            ("id_cupom",       "INT",          "FK → cupons · NN",      "Cupom resgatado"),
            ("codigo_gerado",  "VARCHAR(20)",  "NN · UNIQUE",           "Código único gerado (ex: MESA-ABC123)"),
            ("data_resgate",   "DATETIME",     "NN · DEFAULT NOW()",    "Data do resgate"),
            ("data_uso",       "DATETIME",     "",                      "Data em que foi marcado como usado (NULL = não usado)"),
            ("status",         "ENUM",         "NN · DEFAULT 'ativo'",  "'ativo' | 'usado' | 'expirado'"),
        ],
    },
    "certificados": {
        "desc": "Certificados de horas complementares de voluntariado emitidos para os voluntários.",
        "cols": [
            ("id_certificado",       "INT",          "PK · AUTO_INCREMENT", "Identificador do certificado"),
            ("id_voluntario",        "INT",          "FK → usuarios · NN",  "Voluntário que recebeu o certificado"),
            ("nome_no_certificado",  "VARCHAR(120)", "NN",                  "Nome exibido no certificado"),
            ("total_entregas",       "INT",          "NN",                  "Total de entregas realizadas no período"),
            ("total_horas",          "DECIMAL(6,1)", "NN",                  "Total de horas estimadas (1,5h por entrega)"),
            ("periodo_inicio",       "DATE",         "NN",                  "Início do período coberto"),
            ("periodo_fim",          "DATE",         "NN",                  "Fim do período coberto"),
            ("data_emissao",         "DATETIME",     "NN · DEFAULT NOW()",  "Data e hora da emissão"),
            ("url_pdf",              "VARCHAR(255)", "",                    "Caminho do arquivo PDF gerado"),
        ],
    },
    "notificacoes": {
        "desc": "Notificações enviadas a qualquer usuário do sistema.",
        "cols": [
            ("id_notificacao", "INT",          "PK · AUTO_INCREMENT",  "Identificador da notificação"),
            ("id_usuario",     "INT",          "FK → usuarios · NN",   "Destinatário da notificação"),
            ("icone",          "VARCHAR(10)",  "",                     "Emoji do ícone"),
            ("titulo",         "VARCHAR(120)", "NN",                   "Título curto da notificação"),
            ("detalhe",        "VARCHAR(255)", "",                     "Texto complementar"),
            ("lida",           "TINYINT(1)",  "NN · DEFAULT 0",        "0 = não lida, 1 = lida"),
            ("data_criacao",   "DATETIME",    "NN · DEFAULT NOW()",    "Data de criação"),
        ],
    },
    "mensagens_sac": {
        "desc": "Mensagens de suporte enviadas pelos usuários e respondidas por administradores.",
        "cols": [
            ("id_mensagem",    "INT",      "PK · AUTO_INCREMENT",           "Identificador da mensagem"),
            ("id_usuario",     "INT",      "FK → usuarios · NN",            "Usuário que enviou a mensagem"),
            ("id_atendente",   "INT",      "FK → usuarios",                 "Admin que respondeu (NULL enquanto não atendida)"),
            ("assunto",        "VARCHAR(100)", "NN",                        "Assunto selecionado pelo usuário"),
            ("mensagem",       "TEXT",     "NN",                            "Conteúdo da mensagem"),
            ("resposta",       "TEXT",     "",                              "Resposta do atendente (NULL enquanto aberta)"),
            ("data_envio",     "DATETIME", "NN · DEFAULT NOW()",            "Data de envio"),
            ("data_resposta",  "DATETIME", "",                              "Data da resposta (NULL enquanto pendente)"),
            ("status",         "ENUM",     "NN · DEFAULT 'aberta'",         "'aberta' | 'em_atendimento' | 'respondida' | 'fechada'"),
        ],
    },
}

# ──────────────────────────────────────────────────────────────────────────────
# RELACIONAMENTOS
# ──────────────────────────────────────────────────────────────────────────────
RELATIONSHIPS = [
    {
        "tabela_a": "usuarios",
        "cardinalidade": "1 : 1",
        "tabela_b": "doadores_perfil",
        "campo_fk": "doadores_perfil.id_usuario",
        "descricao": (
            "Cada usuário do tipo 'doador' possui exatamente um perfil de doador, "
            "e cada perfil pertence a exatamente um usuário. "
            "O registro é criado automaticamente junto com o cadastro do doador."
        ),
    },
    {
        "tabela_a": "usuarios",
        "cardinalidade": "1 : 1",
        "tabela_b": "voluntarios_perfil",
        "campo_fk": "voluntarios_perfil.id_usuario",
        "descricao": (
            "Cada usuário do tipo 'voluntario' possui exatamente um perfil de voluntário, "
            "onde ficam armazenados disponibilidade, veículo e saldo de pontos. "
            "Criado junto com o cadastro do voluntário."
        ),
    },
    {
        "tabela_a": "usuarios",
        "cardinalidade": "1 : 1",
        "tabela_b": "familias",
        "campo_fk": "familias.id_usuario",
        "descricao": (
            "Cada usuário do tipo 'beneficiario' está associado a exatamente uma família. "
            "A tabela 'familias' armazena endereço, composição familiar e renda, "
            "complementando os dados de login da tabela 'usuarios'."
        ),
    },
    {
        "tabela_a": "usuarios",
        "cardinalidade": "1 : N",
        "tabela_b": "doacoes",
        "campo_fk": "doacoes.id_doador",
        "descricao": (
            "Um doador pode registrar múltiplas doações ao longo do tempo, "
            "mas cada doação pertence a exatamente um doador. "
            "Sem doador não existe doação."
        ),
    },
    {
        "tabela_a": "usuarios",
        "cardinalidade": "1 : N",
        "tabela_b": "entregas",
        "campo_fk": "entregas.id_voluntario",
        "descricao": (
            "Um voluntário pode realizar múltiplas entregas, "
            "mas cada entrega é atribuída a exatamente um voluntário. "
            "A entrega só existe quando um voluntário a aceita."
        ),
    },
    {
        "tabela_a": "usuarios",
        "cardinalidade": "1 : N",
        "tabela_b": "pontos_historico",
        "campo_fk": "pontos_historico.id_voluntario",
        "descricao": (
            "Um voluntário acumula múltiplos lançamentos de pontos (ganhos e gastos), "
            "mas cada lançamento pertence a exatamente um voluntário. "
            "O saldo em 'voluntarios_perfil.saldo_pontos' é mantido sincronizado com esta tabela."
        ),
    },
    {
        "tabela_a": "usuarios",
        "cardinalidade": "1 : N",
        "tabela_b": "resgates_cupons",
        "campo_fk": "resgates_cupons.id_voluntario",
        "descricao": (
            "Um voluntário pode resgatar múltiplos cupons ao longo do tempo, "
            "mas cada resgate pertence a exatamente um voluntário."
        ),
    },
    {
        "tabela_a": "usuarios",
        "cardinalidade": "1 : N",
        "tabela_b": "certificados",
        "campo_fk": "certificados.id_voluntario",
        "descricao": (
            "Um voluntário pode solicitar múltiplos certificados (ex: um por semestre), "
            "mas cada certificado pertence a exatamente um voluntário."
        ),
    },
    {
        "tabela_a": "usuarios",
        "cardinalidade": "1 : N",
        "tabela_b": "notificacoes",
        "campo_fk": "notificacoes.id_usuario",
        "descricao": (
            "Um usuário pode receber múltiplas notificações. "
            "Vale para todos os tipos de perfil (doador, voluntário, beneficiário, ponto, admin)."
        ),
    },
    {
        "tabela_a": "usuarios",
        "cardinalidade": "1 : N",
        "tabela_b": "mensagens_sac",
        "campo_fk": "mensagens_sac.id_usuario / mensagens_sac.id_atendente",
        "descricao": (
            "Um usuário pode enviar múltiplas mensagens de suporte. "
            "O campo 'id_atendente' também é FK para usuarios, referenciando o admin que respondeu. "
            "Um admin pode responder várias mensagens."
        ),
    },
    {
        "tabela_a": "pontos_coleta",
        "cardinalidade": "1 : N",
        "tabela_b": "estoque_ponto",
        "campo_fk": "estoque_ponto.id_ponto",
        "descricao": (
            "Um ponto de coleta controla o estoque de vários tipos de item. "
            "Cada registro de estoque pertence a um único ponto e a um único tipo de item."
        ),
    },
    {
        "tabela_a": "pontos_coleta",
        "cardinalidade": "1 : N",
        "tabela_b": "doacoes",
        "campo_fk": "doacoes.id_ponto",
        "descricao": (
            "Um ponto de coleta recebe múltiplas doações de diferentes doadores. "
            "Cada doação é destinada a exatamente um ponto."
        ),
    },
    {
        "tabela_a": "pontos_coleta",
        "cardinalidade": "1 : N",
        "tabela_b": "solicitacoes",
        "campo_fk": "solicitacoes.id_ponto",
        "descricao": (
            "Um ponto de coleta recebe e gerencia múltiplas solicitações de famílias. "
            "Cada solicitação é direcionada a exatamente um ponto."
        ),
    },
    {
        "tabela_a": "pontos_coleta",
        "cardinalidade": "1 : N",
        "tabela_b": "familias",
        "campo_fk": "familias.id_ponto_referencia",
        "descricao": (
            "Um ponto de coleta é o ponto de referência de várias famílias da região. "
            "Cada família indica qual ponto fica mais próximo do seu endereço."
        ),
    },
    {
        "tabela_a": "pontos_coleta",
        "cardinalidade": "1 : N",
        "tabela_b": "entregas",
        "campo_fk": "entregas.id_ponto",
        "descricao": (
            "De um mesmo ponto de coleta podem partir múltiplas entregas. "
            "Cada entrega tem um único ponto de retirada."
        ),
    },
    {
        "tabela_a": "itens_catalogo",
        "cardinalidade": "1 : N",
        "tabela_b": "estoque_ponto",
        "campo_fk": "estoque_ponto.id_item",
        "descricao": (
            "Um tipo de item (ex: 'Cesta básica') pode estar no estoque de vários pontos. "
            "Cada registro de estoque refere-se a um único tipo de item."
        ),
    },
    {
        "tabela_a": "itens_catalogo",
        "cardinalidade": "1 : N",
        "tabela_b": "doacoes_itens",
        "campo_fk": "doacoes_itens.id_item",
        "descricao": (
            "Um tipo de item pode aparecer em muitas doações diferentes. "
            "Cada linha de 'doacoes_itens' refere-se a um único tipo de item."
        ),
    },
    {
        "tabela_a": "itens_catalogo",
        "cardinalidade": "1 : N",
        "tabela_b": "solicitacoes_itens",
        "campo_fk": "solicitacoes_itens.id_item",
        "descricao": (
            "Um tipo de item pode aparecer em muitas solicitações diferentes. "
            "Cada linha de 'solicitacoes_itens' refere-se a um único tipo de item."
        ),
    },
    {
        "tabela_a": "doacoes",
        "cardinalidade": "1 : N",
        "tabela_b": "doacoes_itens",
        "campo_fk": "doacoes_itens.id_doacao",
        "descricao": (
            "Uma doação contém um ou mais tipos de itens. "
            "Cada linha de 'doacoes_itens' pertence a exatamente uma doação. "
            "Ao excluir uma doação, suas linhas de itens devem ser excluídas em cascata."
        ),
    },
    {
        "tabela_a": "familias",
        "cardinalidade": "1 : N",
        "tabela_b": "solicitacoes",
        "campo_fk": "solicitacoes.id_familia",
        "descricao": (
            "Uma família pode fazer múltiplas solicitações ao longo do tempo. "
            "Cada solicitação pertence a exatamente uma família."
        ),
    },
    {
        "tabela_a": "solicitacoes",
        "cardinalidade": "1 : 1",
        "tabela_b": "entregas",
        "campo_fk": "entregas.id_solicitacao (UNIQUE)",
        "descricao": (
            "Uma solicitação que precisa de entrega gera exatamente uma entrega. "
            "Uma entrega atende exatamente uma solicitação. "
            "A constraint UNIQUE em 'id_solicitacao' garante esta cardinalidade no banco."
        ),
    },
    {
        "tabela_a": "solicitacoes",
        "cardinalidade": "1 : N",
        "tabela_b": "solicitacoes_itens",
        "campo_fk": "solicitacoes_itens.id_solicitacao",
        "descricao": (
            "Uma solicitação contém um ou mais tipos de itens. "
            "Cada linha de 'solicitacoes_itens' pertence a exatamente uma solicitação."
        ),
    },
    {
        "tabela_a": "entregas",
        "cardinalidade": "1 : N",
        "tabela_b": "pontos_historico",
        "campo_fk": "pontos_historico.id_entrega",
        "descricao": (
            "Uma entrega concluída gera um lançamento de pontos para o voluntário. "
            "O campo 'id_entrega' em 'pontos_historico' pode ser NULL quando o lançamento vem de outro motivo."
        ),
    },
    {
        "tabela_a": "cupons",
        "cardinalidade": "1 : N",
        "tabela_b": "resgates_cupons",
        "campo_fk": "resgates_cupons.id_cupom",
        "descricao": (
            "Um cupom pode ser resgatado por múltiplos voluntários (ou várias vezes pelo mesmo). "
            "Cada resgate refere-se a exatamente um tipo de cupom."
        ),
    },
    {
        "tabela_a": "resgates_cupons",
        "cardinalidade": "1 : N",
        "tabela_b": "pontos_historico",
        "campo_fk": "pontos_historico.id_resgate",
        "descricao": (
            "Um resgate de cupom gera um lançamento de gasto de pontos. "
            "O campo 'id_resgate' em 'pontos_historico' pode ser NULL quando o lançamento vem de entrega."
        ),
    },
]

# ──────────────────────────────────────────────────────────────────────────────
# CORES
# ──────────────────────────────────────────────────────────────────────────────
COR_LARANJA   = "FF7A00"
COR_LARANJA_L = "FFF3E0"
COR_HEADER    = "FF7A00"
COR_PK        = "FFF9C4"
COR_FK        = "E3F2FD"
COR_ZEBRA     = "FAFAFA"
COR_BRANCO    = "FFFFFF"
COR_BORDA     = "E0E0E0"

def make_border():
    thin = Side(border_style="thin", color=COR_BORDA)
    return Border(left=thin, right=thin, top=thin, bottom=thin)

def make_fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)

# ──────────────────────────────────────────────────────────────────────────────
# GERA EXCEL
# ──────────────────────────────────────────────────────────────────────────────
def gerar_excel():
    wb = Workbook()
    wb.remove(wb.active)  # remove a aba padrão

    # Aba índice
    ws_idx = wb.create_sheet("📋 Índice")
    ws_idx.sheet_view.showGridLines = False
    ws_idx.column_dimensions["A"].width = 6
    ws_idx.column_dimensions["B"].width = 28
    ws_idx.column_dimensions["C"].width = 55
    ws_idx.column_dimensions["D"].width = 16

    # Título da aba índice
    ws_idx.merge_cells("A1:D1")
    c = ws_idx["A1"]
    c.value = "Mesa Solidária — Modelo de Banco de Dados"
    c.font = Font(name="Calibri", bold=True, size=16, color="FFFFFF")
    c.fill = make_fill(COR_LARANJA)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws_idx.row_dimensions[1].height = 36

    ws_idx.merge_cells("A2:D2")
    c2 = ws_idx["A2"]
    c2.value = "18 tabelas  ·  MySQL / PostgreSQL"
    c2.font = Font(name="Calibri", italic=True, size=11, color="888888")
    c2.alignment = Alignment(horizontal="center", vertical="center")
    ws_idx.row_dimensions[2].height = 22

    headers_idx = ["#", "Tabela", "Descrição", "Campos a criar"]
    for col, h in enumerate(headers_idx, 1):
        cell = ws_idx.cell(row=4, column=col, value=h)
        cell.font = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
        cell.fill = make_fill("444444")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = make_border()
    ws_idx.row_dimensions[4].height = 22

    for i, (tname, tdata) in enumerate(TABLES.items(), 1):
        row = i + 4
        pk_count  = len([c for c in tdata["cols"] if c[0] and "PK" in c[2]])
        fk_count  = len([c for c in tdata["cols"] if c[0] and "FK" in c[2]])
        reg_count = len([c for c in tdata["cols"] if c[0] and "PK" not in c[2] and "FK" not in c[2]])
        summary = f"PK: {pk_count}  ·  FK: {fk_count}  ·  Dados: {reg_count}"
        values = [i, tname, tdata["desc"], summary]
        for col, val in enumerate(values, 1):
            cell = ws_idx.cell(row=row, column=col, value=val)
            cell.font = Font(name="Calibri", size=10)
            cell.alignment = Alignment(vertical="center", wrap_text=(col == 3))
            cell.border = make_border()
            if col == 1:
                cell.alignment = Alignment(horizontal="center", vertical="center")
            fill_color = COR_ZEBRA if i % 2 == 0 else COR_BRANCO
            cell.fill = make_fill(fill_color)
        ws_idx.row_dimensions[row].height = 28

    # Abas de cada tabela
    for tname, tdata in TABLES.items():
        ws = wb.create_sheet(tname)
        ws.sheet_view.showGridLines = False

        # Larguras: Tipo de campo | Coluna | Tipo SQL | Constraints | Descrição
        widths = [18, 22, 18, 30, 48]
        for col, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(col)].width = w

        # Título
        ws.merge_cells("A1:E1")
        c = ws["A1"]
        c.value = f"  {tname}"
        c.font = Font(name="Calibri", bold=True, size=14, color="FFFFFF")
        c.fill = make_fill(COR_LARANJA)
        c.alignment = Alignment(vertical="center")
        ws.row_dimensions[1].height = 32

        # Descrição
        ws.merge_cells("A2:E2")
        c2 = ws["A2"]
        c2.value = f"  {tdata['desc']}"
        c2.font = Font(name="Calibri", italic=True, size=10, color="555555")
        c2.fill = make_fill(COR_LARANJA_L)
        c2.alignment = Alignment(vertical="center", wrap_text=True)
        ws.row_dimensions[2].height = 36

        # Nota de orientação
        ws.merge_cells("A3:E3")
        cnota = ws["A3"]
        cnota.value = (
            "  ✅  Todos os campos listados abaixo devem ser criados NESTA tabela. "
            "Campos FK (🔗) representam o lado 'filho' do relacionamento — "
            "você cria a coluna aqui, não na tabela referenciada."
        )
        cnota.font = Font(name="Calibri", italic=True, size=9, color="1A5276")
        cnota.fill = make_fill("D6EAF8")
        cnota.alignment = Alignment(vertical="center", wrap_text=True)
        ws.row_dimensions[3].height = 30

        # Cabeçalho colunas
        col_headers = ["Tipo de Campo", "Coluna", "Tipo SQL", "Constraints", "Descrição"]
        for col, h in enumerate(col_headers, 1):
            cell = ws.cell(row=4, column=col, value=h)
            cell.font = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
            cell.fill = make_fill("555555")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = make_border()
        ws.row_dimensions[4].height = 20

        # Linhas das colunas
        for r, (col_name, col_type, col_constraints, col_desc) in enumerate(tdata["cols"], 5):
            if not col_name:
                # Constraint de tabela (ex: UNIQUE composto) — aparece como nota de rodapé
                ws.merge_cells(f"A{r}:E{r}")
                cell = ws.cell(row=r, column=1, value=f"  ⚙  Restrição da tabela: {col_constraints}  —  {col_desc}")
                cell.font = Font(name="Calibri", italic=True, size=9, color="777777")
                cell.fill = make_fill("F5F5F5")
                cell.border = make_border()
                ws.row_dimensions[r].height = 18
                continue

            is_pk = "PK" in col_constraints
            is_fk = "FK" in col_constraints
            fill_color = COR_PK if is_pk else (COR_FK if is_fk else (COR_ZEBRA if r % 2 == 0 else COR_BRANCO))

            if is_pk:
                tipo_label = "🔑 Chave Primária"
                tipo_color = "B7950B"
            elif is_fk:
                tipo_label = "🔗 Chave Estrangeira"
                tipo_color = "1A5276"
            else:
                tipo_label = "📊 Dado"
                tipo_color = "2C3E50"

            row_vals = [tipo_label, col_name, col_type, col_constraints, col_desc]
            for col, val in enumerate(row_vals, 1):
                cell = ws.cell(row=r, column=col, value=val)
                cell.font = Font(
                    name="Calibri",
                    size=10,
                    bold=(col == 2 and is_pk),
                    color=(
                        tipo_color if col == 1
                        else ("CC0000" if is_pk and col == 2
                              else ("0066CC" if is_fk and col == 4
                                    else "222222"))
                    ),
                )
                cell.fill = make_fill(fill_color)
                cell.alignment = Alignment(vertical="center", wrap_text=(col in (4, 5)))
                cell.border = make_border()
            ws.row_dimensions[r].height = 22

    # Legenda de cores
    ws_leg = wb.create_sheet("🎨 Legenda")
    ws_leg.sheet_view.showGridLines = False
    ws_leg.column_dimensions["A"].width = 28
    ws_leg.column_dimensions["B"].width = 20
    ws_leg.column_dimensions["C"].width = 60

    ws_leg.merge_cells("A1:C1")
    c = ws_leg["A1"]
    c.value = "Legenda — Como ler a planilha"
    c.font = Font(name="Calibri", bold=True, size=13, color="FFFFFF")
    c.fill = make_fill(COR_LARANJA)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws_leg.row_dimensions[1].height = 28

    # Cabeçalho legenda
    for col, h in enumerate(["Cor da linha", "Tipo de Campo", "O que significa / O que você deve fazer"], 1):
        cell = ws_leg.cell(row=2, column=col, value=h)
        cell.font = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
        cell.fill = make_fill("444444")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = make_border()
    ws_leg.row_dimensions[2].height = 20

    legenda_items = [
        (COR_PK,    "🔑 Chave Primária",    "É o identificador único de cada linha. Crie SEMPRE esta coluna na tabela. Geralmente é AUTO_INCREMENT (gerado automaticamente pelo banco)."),
        (COR_FK,    "🔗 Chave Estrangeira", "Liga esta tabela a outra. Crie esta coluna SOMENTE AQUI. A tabela referenciada NÃO precisa de uma coluna de volta — o relacionamento é mantido por este campo."),
        (COR_ZEBRA, "📊 Dado (par)",        "Campo de dado normal. Crie esta coluna normalmente na tabela."),
        (COR_BRANCO,"📊 Dado (ímpar)",      "Campo de dado normal. Crie esta coluna normalmente na tabela."),
        ("D6EAF8",  "ℹ️  Nota azul",        "Orientação geral no topo de cada aba. Apenas informativo."),
        ("F5F5F5",  "⚙  Restrição",         "Não é uma coluna — é uma regra extra da tabela (ex: UNIQUE em dois campos juntos). Adicione como CONSTRAINT no CREATE TABLE."),
    ]
    for i, (cor, tipo, texto) in enumerate(legenda_items, 3):
        cell_cor = ws_leg.cell(row=i, column=1, value="")
        cell_cor.fill = make_fill(cor)
        cell_cor.border = make_border()
        cell_tipo = ws_leg.cell(row=i, column=2, value=tipo)
        cell_tipo.font = Font(name="Calibri", bold=True, size=10)
        cell_tipo.fill = make_fill(cor)
        cell_tipo.alignment = Alignment(vertical="center")
        cell_tipo.border = make_border()
        c2 = ws_leg.cell(row=i, column=3, value=texto)
        c2.font = Font(name="Calibri", size=10)
        c2.alignment = Alignment(vertical="center", wrap_text=True)
        c2.border = make_border()
        ws_leg.row_dimensions[i].height = 36

    path = os.path.join(OUTPUT_DIR, "banco_de_dados.xlsx")
    wb.save(path)
    print(f"✅ Excel gerado: {path}")

# ──────────────────────────────────────────────────────────────────────────────
# GERA WORD
# ──────────────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Define cor de fundo de uma célula de tabela Word."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color_hex="FF7A00"):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return p

def gerar_word():
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ── Capa ─────────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("Mesa Solidária")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x7A, 0x00)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_sub.add_run("Modelo de Banco de Dados — Relacionamentos e Cardinalidade")
    run2.font.name = "Calibri"
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # ── Introdução ────────────────────────────────────────────────────────────
    add_heading(doc, "1. Visão Geral", level=1)
    intro = doc.add_paragraph(
        "O banco de dados do Mesa Solidária é composto por 18 tabelas organizadas em torno "
        "de cinco perfis de usuário: Doador, Voluntário, Beneficiário, Ponto de Coleta e Administrador. "
        "Todas as contas de acesso são armazenadas na tabela central usuarios, enquanto tabelas "
        "de perfil complementam os dados específicos de cada tipo. "
        "As operações do sistema (doações, solicitações, entregas, pontos e cupons) "
        "são rastreadas por tabelas transacionais, garantindo auditabilidade completa."
    )
    intro.paragraph_format.space_after = Pt(8)

    # Tabela resumo
    add_heading(doc, "2. Resumo das Tabelas", level=1)
    tbl_res = doc.add_table(rows=1, cols=3)
    tbl_res.style = "Table Grid"
    tbl_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tbl_res.rows[0].cells
    for i, h in enumerate(["Tabela", "Finalidade", "Tipo"]):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        set_cell_bg(hdr_cells[i], "FF7A00")
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    tipos_tabela = {
        "usuarios":            ("Todos os perfis de acesso",                        "Central"),
        "doadores_perfil":     ("Dados extras do doador",                           "Perfil"),
        "voluntarios_perfil":  ("Dados extras + pontos do voluntário",              "Perfil"),
        "pontos_coleta":       ("Pontos físicos de coleta e distribuição",          "Operacional"),
        "familias":            ("Dados cadastrais das famílias beneficiárias",      "Perfil"),
        "itens_catalogo":      ("Catálogo de tipos de itens doáveis",               "Referência"),
        "estoque_ponto":       ("Estoque atual por item e por ponto",               "Operacional"),
        "doacoes":             ("Registros de doações",                             "Transacional"),
        "doacoes_itens":       ("Itens de cada doação",                             "Junção"),
        "solicitacoes":        ("Pedidos das famílias ao ponto",                    "Transacional"),
        "solicitacoes_itens":  ("Itens de cada solicitação",                        "Junção"),
        "entregas":            ("Entregas aceitas por voluntários",                 "Transacional"),
        "pontos_historico":    ("Lançamentos de pontos dos voluntários",            "Auditoria"),
        "cupons":              ("Catálogo de recompensas",                          "Referência"),
        "resgates_cupons":     ("Cupons resgatados com código único",               "Transacional"),
        "certificados":        ("Certificados de horas emitidos",                   "Transacional"),
        "notificacoes":        ("Notificações por usuário",                         "Suporte"),
        "mensagens_sac":       ("Mensagens de suporte ao cliente",                  "Suporte"),
    }
    cores_tipo = {
        "Central":      "FFF3E0",
        "Perfil":       "E3F2FD",
        "Operacional":  "E8F5E9",
        "Referência":   "F3E5F5",
        "Transacional": "FFFFFF",
        "Junção":       "FFF9C4",
        "Auditoria":    "FCE4EC",
        "Suporte":      "F5F5F5",
    }
    for tname, (tdesc, ttipo) in tipos_tabela.items():
        row_cells = tbl_res.add_row().cells
        row_cells[0].text = tname
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = tdesc
        row_cells[2].text = ttipo
        cor = cores_tipo.get(ttipo, "FFFFFF")
        for cell in row_cells:
            set_cell_bg(cell, cor)

    doc.add_paragraph()

    # ── Relacionamentos ───────────────────────────────────────────────────────
    add_heading(doc, "3. Relacionamentos e Cardinalidade", level=1)

    intro2 = doc.add_paragraph(
        "A seguir, cada relacionamento é descrito com sua cardinalidade e campo de ligação. "
        "A notação utilizada é  A (1) ──── (N) B,  onde A é o lado '1' (pai) "
        "e B é o lado 'muitos' (filho), salvo quando indicado 1:1."
    )
    intro2.paragraph_format.space_after = Pt(10)

    # Agrupa por tabela pai
    grupos = {}
    for rel in RELATIONSHIPS:
        grupos.setdefault(rel["tabela_a"], []).append(rel)

    for tabela_pai, rels in grupos.items():
        add_heading(doc, f"3.{list(grupos.keys()).index(tabela_pai)+1}  {tabela_pai}", level=2, color_hex="1565C0")
        for rel in rels:
            # Linha de cardinalidade
            p_card = doc.add_paragraph()
            p_card.paragraph_format.left_indent = Cm(0.8)
            r1 = p_card.add_run(f"{rel['tabela_a']}")
            r1.font.bold = True
            r1.font.color.rgb = RGBColor(0xFF, 0x7A, 0x00)
            r2 = p_card.add_run(f"  ({rel['cardinalidade']})  ")
            r2.font.bold = True
            r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            r3 = p_card.add_run(rel["tabela_b"])
            r3.font.bold = True
            r3.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

            # FK
            p_fk = doc.add_paragraph()
            p_fk.paragraph_format.left_indent = Cm(1.4)
            rfk = p_fk.add_run(f"Campo: {rel['campo_fk']}")
            rfk.font.italic = True
            rfk.font.size = Pt(9.5)
            rfk.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

            # Descrição
            p_desc = doc.add_paragraph()
            p_desc.paragraph_format.left_indent = Cm(1.4)
            p_desc.paragraph_format.space_after = Pt(10)
            rdesc = p_desc.add_run(rel["descricao"])
            rdesc.font.size = Pt(10)

    # ── Cardinalidades únicas ─────────────────────────────────────────────────
    add_heading(doc, "4. Tabela consolidada de relacionamentos", level=1)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Cm(4), Cm(2.5), Cm(4), Cm(7)]
    for i, (hdr, w) in enumerate(zip(["Tabela Pai", "Cardinalidade", "Tabela Filho", "Campo FK"], col_widths)):
        cell = tbl.rows[0].cells[i]
        cell.width = w
        cell.text = hdr
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "FF7A00")

    card_cores = {"1 : 1": "FFF9C4", "1 : N": "FFFFFF", "N : M": "FCE4EC"}
    for i, rel in enumerate(RELATIONSHIPS):
        row_cells = tbl.add_row().cells
        row_cells[0].text = rel["tabela_a"]
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        c_card = row_cells[1]
        c_card.text = rel["cardinalidade"]
        c_card.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].text = rel["tabela_b"]
        row_cells[3].text = rel["campo_fk"]
        row_cells[3].paragraphs[0].runs[0].font.size = Pt(9)
        cor = card_cores.get(rel["cardinalidade"], "FFFFFF")
        if i % 2 == 0:
            cor = cor if cor != "FFFFFF" else "FAFAFA"
        set_cell_bg(c_card, cor)

    doc.add_paragraph()

    # ── Regras de negócio ─────────────────────────────────────────────────────
    add_heading(doc, "5. Regras de negócio importantes", level=1)
    regras = [
        ("Cascade Delete", "Ao excluir uma doação, os registros em doacoes_itens devem ser excluídos em cascata (ON DELETE CASCADE)."),
        ("Unique Constraint", "O par (id_ponto, id_item) em estoque_ponto é UNIQUE, impedindo estoque duplicado do mesmo item no mesmo ponto."),
        ("Solicitação → Entrega", "Uma solicitação só pode ter UMA entrega vinculada (UNIQUE em entregas.id_solicitacao). Para modalidade 'retirada', não é criada entrega."),
        ("Saldo de pontos", "O campo voluntarios_perfil.saldo_pontos é atualizado pela aplicação sempre que um lançamento é inserido em pontos_historico."),
        ("Aprovação de beneficiário", "Novos beneficiários entram com status 'pendente' em familias. Só após aprovação do admin (status → 'ativa') podem receber solicitações."),
        ("Código de cupom", "O campo resgates_cupons.codigo_gerado é UNIQUE para garantir que dois resgates nunca tenham o mesmo código."),
        ("Self-join em mensagens_sac", "A tabela mensagens_sac possui dois FKs para usuarios: id_usuario (quem enviou) e id_atendente (admin que respondeu)."),
    ]
    for titulo, texto in regras:
        p = doc.add_paragraph(style="List Bullet")
        r_t = p.add_run(f"{titulo}: ")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(0xFF, 0x7A, 0x00)
        p.add_run(texto)

    doc.add_paragraph()
    # Rodapé
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = p_footer.add_run("Mesa Solidária  ·  Documentação de banco de dados  ·  Gerado automaticamente")
    rf.font.size = Pt(8)
    rf.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    path = os.path.join(OUTPUT_DIR, "relacionamentos_db.docx")
    doc.save(path)
    print(f"✅ Word gerado: {path}")

# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    if "--word-only" in sys.argv:
        gerar_word()
    elif "--excel-only" in sys.argv:
        gerar_excel()
    else:
        gerar_excel()
        gerar_word()
    print("\n🎉 Documentação gerada com sucesso!")
    print(f"   📊 banco_de_dados.xlsx")
    print(f"   📄 relacionamentos_db.docx")
